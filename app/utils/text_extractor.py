"""
Text extraction utilities for PDF and DOCX files
"""
import io
from typing import Optional
from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager
from docx import Document
from fastapi import HTTPException


class TextExtractor:
    """Utility class for extracting text from various file formats"""
    
    @staticmethod
    def extract_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file content
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text as string
            
        Raises:
            HTTPException: If PDF processing fails
        """
        try:
            # Create a string buffer to capture the extracted text
            text_buffer = io.StringIO()
            
            # Create PDF resource manager and layout parameters
            resource_manager = PDFResourceManager()
            laparams = LAParams(
                char_margin=2.0,
                line_margin=0.5,
                word_margin=0.1,
                boxes_flow=0.5,
                all_texts=False
            )
            
            # Extract text to the buffer
            extract_text_to_fp(
                io.BytesIO(file_content),
                text_buffer,
                laparams=laparams,
                output_type='text',
                codec='utf-8'
            )
            
            # Get the extracted text
            extracted_text = text_buffer.getvalue()
            text_buffer.close()
            
            # Clean up the text
            cleaned_text = TextExtractor._clean_text(extracted_text)
            
            if not cleaned_text or len(cleaned_text.strip()) < 10:
                raise HTTPException(
                    status_code=400,
                    detail="Could not extract meaningful text from PDF. The file might be corrupted, password-protected, or contain only images."
                )
                
            return cleaned_text
            
        except Exception as e:
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(
                status_code=400,
                detail=f"Failed to process PDF file: {str(e)}"
            )
    
    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file content
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text as string
            
        Raises:
            HTTPException: If DOCX processing fails
        """
        try:
            # Create a BytesIO object from the file content
            doc_buffer = io.BytesIO(file_content)
            
            # Load the document
            doc = Document(doc_buffer)
            
            # Extract text from all paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        table_texts.append(" | ".join(row_texts))
            
            # Combine all text
            all_text = paragraphs + table_texts
            extracted_text = "\n".join(all_text)
            
            # Clean up the text
            cleaned_text = TextExtractor._clean_text(extracted_text)
            
            if not cleaned_text or len(cleaned_text.strip()) < 10:
                raise HTTPException(
                    status_code=400,
                    detail="Could not extract meaningful text from DOCX file. The file might be corrupted or contain only images."
                )
                
            return cleaned_text
            
        except Exception as e:
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(
                status_code=400,
                detail=f"Failed to process DOCX file: {str(e)}"
            )
    
    @staticmethod
    def extract_text(file_content: bytes, file_extension: str) -> str:
        """
        Extract text from file based on its extension
        
        Args:
            file_content: File content as bytes
            file_extension: File extension (e.g., '.pdf', '.docx')
            
        Returns:
            Extracted text as string
            
        Raises:
            HTTPException: If file format is not supported or processing fails
        """
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            return TextExtractor.extract_from_pdf(file_content)
        elif file_extension in ['.docx', '.doc']:
            return TextExtractor.extract_from_docx(file_content)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx"
            )
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Join lines with proper spacing
        cleaned_text = '\n'.join(lines)
        
        # Remove excessive spaces between words
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        # Remove special characters that might be artifacts from PDF extraction
        cleaned_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', cleaned_text)
        
        return cleaned_text.strip()
